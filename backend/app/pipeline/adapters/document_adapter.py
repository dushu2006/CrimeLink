"""FIR / surveillance / intelligence report adapter (PRD 7).

Handles PDF (with a native text layer), DOCX and plain text.  Page and line
offsets are preserved so ``text_span`` evidence pointers resolve to the exact
sentence in the original document.

**OCR is not bundled.**  A PDF that carries no extractable text (a scan or a
photograph of an FIR) fails with an explicit reason — "OCR is required" — and is
quarantined, rather than silently yielding an empty graph.  Adding OCR means
installing the ``tesseract`` binary plus ``hi+eng`` traineddata on the worker
image and calling it from :meth:`TextDocumentAdapter._parse_pdf`; the failure
path and the quarantine UI are already in place to receive it.
"""

from __future__ import annotations

import io
import shutil
import subprocess
import tempfile
from pathlib import Path

from app.domain.enums import DocumentType
from app.domain.models import Block, NormalizedDocument, OriginRef
from app.errors import PipelineError
from app.logging import get_logger
from app.pipeline.adapters.protocol import (
    DocumentMeta,
    detect_language,
    text_blocks_from_text,
)

log = get_logger("crimelink.adapter.document")

MIN_CHARS_PER_PAGE = 20


class TextDocumentAdapter:
    """PDF / DOCX / TXT → normalised document."""

    document_type = DocumentType.FIR

    def parse(self, raw: bytes, doc_meta: DocumentMeta) -> NormalizedDocument:
        suffix = Path(doc_meta.filename or "").suffix.lower()
        pages: list[tuple[int, str]] = []
        warnings: list[str] = []

        if suffix == ".pdf" or raw[:4] == b"%PDF":
            pages, warnings = self._parse_pdf(raw)
        elif suffix in (".docx", ".doc"):
            pages.append((1, self._parse_docx(raw)))
        else:
            try:
                pages.append((1, raw.decode("utf-8")))
            except UnicodeDecodeError:
                pages.append((1, raw.decode("latin-1", errors="replace")))
                warnings.append("File was not valid UTF-8; decoded as Latin-1.")

        text = "\n".join(page_text for _, page_text in pages)
        if not text.strip():
            raise PipelineError(
                "No readable text could be extracted from this file. "
                "If it is a scanned document, OCR is required."
            )

        # A report is ingested verbatim, so the uploaded file is its own origin.
        # Recording it per block lets a finding open the exact line (or PDF
        # page) of the real document instead of only the derived rendering.
        origin_file = (doc_meta.extra or {}).get("relative_path") or doc_meta.filename or ""

        blocks: list[Block] = []
        cursor = 0
        for page_number, page_text in pages:
            normalised_page = page_text.replace("\r\n", "\n")
            for block in text_blocks_from_text(normalised_page, page=page_number):
                block.offset += cursor
                if origin_file:
                    block.origin = OriginRef(
                        file=origin_file,
                        row=block.line,
                        fields=["text"],
                        values={
                            "page": page_number,
                            "line": block.line,
                            # A short excerpt keeps the reference legible in
                            # listings without duplicating the document.
                            "text": block.text[:200],
                        },
                    )
                blocks.append(block)
            cursor += len(normalised_page) + 1

        return NormalizedDocument(
            doc_id=doc_meta.doc_id,
            case_id=doc_meta.case_id,
            doc_type=doc_meta.document_type.value,
            language=detect_language(text, doc_meta.language_hint),
            source_confidence=doc_meta.source_confidence,
            blocks=blocks,
            parse_warnings=warnings,
            metadata={"filename": doc_meta.filename, "pages": len(pages)},
        )

    # ------------------------------------------------------------------- pdf
    def _parse_pdf(self, raw: bytes) -> tuple[list[tuple[int, str]], list[str]]:
        warnings: list[str] = []
        try:
            from pypdf import PdfReader
            from pypdf.errors import PdfReadError
        except ImportError as exc:  # pragma: no cover
            raise PipelineError("PDF support is not installed.") from exc

        try:
            reader = PdfReader(io.BytesIO(raw))
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    raise PipelineError(
                        "This PDF is password-protected and could not be opened."
                    )
            pages: list[tuple[int, str]] = []
            for index, page in enumerate(reader.pages, start=1):
                try:
                    pages.append((index, page.extract_text() or ""))
                except Exception:
                    pages.append((index, ""))
                    warnings.append(f"Text could not be extracted from page {index}.")
        except PipelineError:
            raise
        except Exception as exc:
            raise PipelineError(f"This file is not a readable PDF ({type(exc).__name__}).") from exc

        usable = sum(1 for _, text in pages if len(text.strip()) >= MIN_CHARS_PER_PAGE)
        if usable == 0 and pages:
            ocr_pages, ocr_warnings = self._ocr_pdf(raw, len(pages))
            warnings.extend(ocr_warnings)
            if ocr_pages:
                return ocr_pages, warnings
            raise PipelineError(
                "This PDF appears to be a scan (no selectable text) and OCR is not "
                "available on this host. Install tesseract with the 'hin+eng' "
                "language data, or upload a text-based PDF."
            )
        return pages, warnings

    def _ocr_pdf(self, raw: bytes, page_count: int) -> tuple[list[tuple[int, str]], list[str]]:
        """Best-effort OCR.  Returns empty output when tesseract is unavailable."""
        warnings: list[str] = []
        if shutil.which("tesseract") is None:
            return [], warnings
        try:  # pragma: no cover - requires an OCR toolchain
            import pypdfium2  # type: ignore
        except ImportError:
            warnings.append(
                "Tesseract is installed but the PDF rasteriser (pypdfium2) is not, "
                "so OCR was skipped."
            )
            return [], warnings
        pages: list[tuple[int, str]] = []
        try:  # pragma: no cover
            pdf = pypdfium2.PdfDocument(io.BytesIO(raw))
            with tempfile.TemporaryDirectory() as tmp:
                for index in range(len(pdf)):
                    image_path = Path(tmp) / f"page_{index + 1}.png"
                    pdf[index].render(scale=2).to_pil().save(image_path)
                    result = subprocess.run(
                        ["tesseract", str(image_path), "stdout", "-l", "hin+eng"],
                        capture_output=True,
                        timeout=180,
                    )
                    pages.append((index + 1, result.stdout.decode("utf-8", errors="replace")))
        except Exception as exc:  # pragma: no cover
            warnings.append(f"OCR failed: {type(exc).__name__}")
            return [], warnings
        warnings.append("Document was OCR-processed; text accuracy may be reduced.")
        return pages, warnings

    # ------------------------------------------------------------------ docx
    @staticmethod
    def _parse_docx(raw: bytes) -> str:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:  # pragma: no cover
            raise PipelineError("DOCX support is not installed.") from exc
        try:
            document = Document(io.BytesIO(raw))
        except Exception as exc:
            raise PipelineError(f"This file is not a readable DOCX ({type(exc).__name__}).") from exc
        paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs)


class IntelReportAdapter(TextDocumentAdapter):
    """Intelligence reports — same parsing, different provenance classification."""

    document_type = DocumentType.INTEL


class AnonymousTipAdapter(TextDocumentAdapter):
    """Free-text officer entry for an anonymous tip (PRD 7).

    Tips are always ``ANONYMOUS_TIP`` confidence, which makes every entity
    derived from them ``staging=true`` so they cannot enter the main case graph
    until an investigator promotes them (DPDP / legal-review gate).
    """

    document_type = DocumentType.INTEL

    def parse(self, raw: bytes, doc_meta: DocumentMeta) -> NormalizedDocument:
        document = super().parse(raw, doc_meta)
        document.metadata["staging"] = True
        document.metadata["legal_review_required"] = True
        document.parse_warnings.append(
            "Anonymous tip: extracted entities are staged and require investigator "
            "promotion before they join the case graph."
        )
        return document
